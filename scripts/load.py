''' File for writing the SIMT file from the extracted data. '''
import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from utils.file_utils import read_yaml_file
# from scripts.transform import generate_simt_line
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import pandas as pd

def generate_simt_line(fields, data_row = None) :

    line = [' '] * 500 #fixed width

    for field in fields : 
        for name, props in field.items() :
            pos = props["starting position"]
            length = props["longueur"]
            obligatoire = props["obligatoire"]
            type = props["type"]
            default = props["default"]

            if data_row is not None and name in data_row and pd.notna(data_row[name]):
                value = str(data_row[name])
            else :
                # if obligatoire : 
                    if type == 'text':
                        if default != "":
                            value = str(default)
                        else:
                            value = " " * length
                    elif type == 'integer':
                        value = str(default)
                    elif type == 'date' or type == 'time':
                        if default == "today":
                            value = datetime.today().strftime("%Y%m%d")
                        elif default == "now":
                            value = datetime.today().strftime("%H%M%S") 
            if len(value) < length :
                if type == "integer": 
                    value = value.zfill(length)
                else:
                    value = value.ljust(length)

            line[pos:pos+length] = list(value) 

    return ''.join(line)



def generate_simt_file(yaml_path, df, extension="txt", month=None, year=None):
    from datetime import datetime

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if extension not in ["txt", "docx", "asc"]:
        raise ValueError("Extension must be either 'txt', 'docx' or 'asc")
    
    project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
    output_dir = os.path.join(project_root, 'data', 'fee_payouts')
    os.makedirs(output_dir, exist_ok=True) 

    date_emission_dt = pd.to_datetime(df["date_emission"])
    if month is not None and year is not None:
        selected_month = month
        selected_year = year
    else:
        latest_date = date_emission_dt.max()
        selected_month = latest_date.month
        selected_year = latest_date.year

    df_filtered = df[(date_emission_dt.dt.month == selected_month) & (date_emission_dt.dt.year == selected_year)].copy()

    if "reference_remise" in df_filtered.columns and not df_filtered.empty:
        reference_remise = str(df_filtered['reference_remise'].iloc[0])
    else:
        reference_remise = "NOREF"
    output_path = os.path.join(output_dir, f'fee_payouts_{reference_remise}_{timestamp}.{extension}')
    
    # Check if the YAML file exists and read its structure
    if not os.path.exists(yaml_path):
        raise FileNotFoundError(f"YAML file not found: {yaml_path}")
    else:
        structure = read_yaml_file(yaml_path)["integration_file_structure"]
    
    lines = []
    # Header : 
    if not df_filtered.empty:
        header_row = df_filtered.iloc[0]
        lines.append(generate_simt_line(structure["Header"]["Fields"], data_row=header_row))

    # Detail
    nb_virements = 0
    montant_total = 0
    for _, row in df_filtered.iterrows():
        lines.append(generate_simt_line(structure["Detail"]["Fields"], data_row=row))
        nb_virements += 1
        montant_total += float(row["montant"])

    # Footer
    footer_fields = structure["Footer"]["Fields"]
    for field in footer_fields:
        if "nombre_total_virements" in field:
            field["nombre_total_virements"]["default"] = nb_virements
        if "montant_total_virements" in field:
            field["montant_total_virements"]["default"] = montant_total

    lines.append(generate_simt_line(structure["Footer"]["Fields"]))

    # Write to file
    if extension == "txt" or extension == "asc":
        with open(output_path, "w", encoding="utf-8") as f:
            for line in lines:
                f.write(line + "\n")
    elif extension == "docx":
        doc = Document()
        heading = doc.add_heading("SIMT File\n", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(20)
        heading_run.bold = True

        for line in lines:
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.runs[0]
            run.font.size = Pt(12)

        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.25)
            section.right_margin = Inches(0.25)

        doc.save(output_path)

    print(f"Fichier généré : {output_path}")